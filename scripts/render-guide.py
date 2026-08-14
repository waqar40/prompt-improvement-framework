# ai-gen - Render a per-user prompting guide (guides/<user>.json) into cognitively friendly views.
"""Usage: python scripts/render-guide.py guides/<user>.json [--pdf] [--docx] [--md]

With no format flag, emits all three. Output files sit next to the JSON
(<user>.pdf / <user>.docx / <user>.md). One structured JSON in, three human-readable
views out - markdown, PDF, and Word are all just views of the same data.

Each reviewed prompt renders two tables:
  1. Rubric scorecard - one row per rubric dimension (D1-D10 design, E1-E4 evaluability),
     showing which was Met / Partial / Missing / n-a, with the evidence.
  2. Transformation table - one row per gap: what you wrote -> the best-practice rewrite,
     and the prompt-engineering principle it teaches.
"""
from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path
from xml.sax.saxutils import escape

# --------------------------------------------------------------------------- model
# Single source of truth for dimension names + weights (mirrors prompt-critic rubric).
DIMS = [
    ("D1", "Clarity & explicitness", "design", 3),
    ("D2", "Specificity & constraints", "design", 3),
    ("D3", "Output format & length", "design", 2),
    ("D4", "Context & motivation", "design", 2),
    ("D5", "Grounding / reference", "design", 2),
    ("D6", "Examples (show-not-tell)", "design", 2),
    ("D7", "Positive framing", "design", 1),
    ("D8", "Uncertainty handling", "design", 1),
    ("D9", "Decomposition fit", "design", 2),
    ("D10", "Structural economy", "design", 2),
    ("E1", "Success is defined", "eval", 3),
    ("E2", "Criteria are measurable", "eval", 3),
    ("E3", "Multidimensional coverage", "eval", 2),
    ("E4", "Failure modes anticipated", "eval", 2),
]
DIM_NAME = {d[0]: d[1] for d in DIMS}
DIM_WEIGHT = {d[0]: d[3] for d in DIMS}
PTS = {"met": 1.0, "partial": 0.5, "gap": 0.0}

# verdict -> (label, accent hex, light-bg hex)
VERDICT = {
    "met": ("Met", "#1a7f37", "#e8f5ec"),
    "partial": ("Partial", "#9a6700", "#fdf5e6"),
    "gap": ("Missing", "#b42318", "#fdeceb"),
    "na": ("n/a", "#8a8f96", "#f0f1f3"),
}
# verdict -> plain-text marker for markdown
MD_MARK = {"met": "Met", "partial": "Partial", "gap": "Missing", "na": "-"}

# Band palette: (accent, light background) as hex.
BANDS = {
    "excellent": ("#1a7f37", "#e8f5ec"),
    "good": ("#9a6700", "#fdf5e6"),
    "bad": ("#b42318", "#fdeceb"),
}
BAND_LABEL = {"excellent": "EXCELLENT", "good": "GOOD", "bad": "BAD"}


def load(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def rollup_score(rubric: list[dict]) -> int:
    """Recompute the weighted score from the rubric rows (applicable rows only)."""
    wt = aw = 0.0
    for r in rubric:
        if r["verdict"] == "na":
            continue
        w = DIM_WEIGHT[r["id"]]
        wt += w
        aw += w * PTS[r["verdict"]]
    return round(100 * aw / wt) if wt else 0


def dim_row(ex_rubric: list[dict]) -> dict:
    return {r["id"]: r for r in ex_rubric}


def verify(g: dict) -> list[str]:
    """Return warnings where a rubric roll-up disagrees with the stored score."""
    warns = []
    for sec in g["sections"]:
        for ex in sec["examples"]:
            got = rollup_score(ex["rubric"])
            if got != ex["score"]:
                warns.append(f"  ! {ex['label']!r}: stored {ex['score']} vs rubric roll-up {got}")
    return warns


# --------------------------------------------------------------------------- PDF
def render_pdf(g: dict, out: Path) -> None:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        KeepTogether,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    CONTENT_W = 170 * mm  # A4 width 210 - 20 - 20 margins
    HEX = colors.HexColor

    ss = getSampleStyleSheet()
    body = ParagraphStyle("body", parent=ss["BodyText"], fontSize=9.5, leading=13)
    h1 = ParagraphStyle("h1", parent=ss["Title"], fontSize=20, spaceAfter=2)
    sub = ParagraphStyle("sub", parent=body, textColor=HEX("#666666"), fontSize=9)
    prompt = ParagraphStyle("prompt", parent=body, fontName="Helvetica-Oblique",
                            fontSize=10, leading=14, textColor=HEX("#111111"))
    fixstyle = ParagraphStyle("fix", parent=body, fontSize=9.5, leading=13,
                              textColor=HEX("#0b3d91"))
    cell = ParagraphStyle("cell", parent=body, fontSize=8, leading=10)
    cellb = ParagraphStyle("cellb", parent=cell, fontName="Helvetica-Bold")
    cellw = ParagraphStyle("cellw", parent=cell, textColor=colors.white,
                           fontName="Helvetica-Bold", alignment=1)
    thstyle = ParagraphStyle("th", parent=cell, fontName="Helvetica-Bold",
                             textColor=colors.white)

    def RT(markup, style=body):
        """Rich text: caller supplies markup; dynamic pieces must be pre-escaped."""
        return Paragraph(markup, style)

    def PT(text, style=body):
        """Plain text: escaped."""
        return Paragraph(escape(str(text)), style)

    def boxed(flowables, bg, border, width=CONTENT_W, pad=8):
        t = Table([[flowables]], colWidths=[width])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), HEX(bg)),
            ("BOX", (0, 0), (-1, -1), 0.75, HEX(border)),
            ("LEFTPADDING", (0, 0), (-1, -1), pad),
            ("RIGHTPADDING", (0, 0), (-1, -1), pad),
            ("TOPPADDING", (0, 0), (-1, -1), 7),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
        ]))
        return t

    def rubric_table(ex):
        by = dim_row(ex["rubric"])
        data = [[RT("<b>Rubric dimension</b>", thstyle),
                 RT("<b>Status</b>", thstyle),
                 RT("<b>Why (evidence)</b>", thstyle)]]
        styles = [
            ("BACKGROUND", (0, 0), (-1, 0), HEX("#334155")),
            ("GRID", (0, 0), (-1, -1), 0.5, HEX("#cfd4da")),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("LEFTPADDING", (0, 0), (-1, -1), 5),
            ("RIGHTPADDING", (0, 0), (-1, -1), 5),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ]
        r = 1
        for layer, title in (("design", "LAYER 1 - DESIGN  (does the prompt drive good behaviour?)"),
                             ("eval", "LAYER 2 - EVALUABILITY  (can success be measured?)")):
            data.append([RT(f"<b>{escape(title)}</b>", cellb), "", ""])
            styles += [("SPAN", (0, r), (-1, r)),
                       ("BACKGROUND", (0, r), (-1, r), HEX("#eef1f4"))]
            r += 1
            for did, name, dl, _w in DIMS:
                if dl != layer:
                    continue
                row = by.get(did, {"verdict": "na", "evidence": ""})
                vlabel, acc, lt = VERDICT[row["verdict"]]
                data.append([
                    RT(f"<b>{did}</b>&nbsp; {escape(name)}", cell),
                    RT(escape(vlabel), cellw),
                    PT(row.get("evidence", ""), cell),
                ])
                styles.append(("BACKGROUND", (1, r), (1, r), HEX(acc)))
                r += 1
        t = Table(data, colWidths=[58 * mm, 22 * mm, 90 * mm], repeatRows=1)
        t.setStyle(TableStyle(styles))
        return t

    def transform_table(ex):
        rows = ex.get("transformation") or []
        if not rows:
            return None
        data = [[RT("<b>Rubric</b>", thstyle), RT("<b>You wrote</b>", thstyle),
                 RT("<b>Best-practice rewrite</b>", thstyle), RT("<b>Principle</b>", thstyle)]]
        styles = [
            ("BACKGROUND", (0, 0), (-1, 0), HEX("#0b3d91")),
            ("GRID", (0, 0), (-1, -1), 0.5, HEX("#b9cdf0")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 5),
            ("RIGHTPADDING", (0, 0), (-1, -1), 5),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, HEX("#f4f8ff")]),
        ]
        for row in rows:
            did = row.get("id", "")
            tag = f"<b>{escape(did)}</b>" + (f"<br/><font size=6 color='#888888'>{escape(DIM_NAME.get(did, ''))}</font>" if did in DIM_NAME else "")
            data.append([
                RT(tag, cell),
                PT(row.get("before", ""), cell),
                RT(escape(row.get("after", "")), ParagraphStyle("aft", parent=cell, textColor=HEX("#0b3d91"))),
                PT(row.get("principle", ""), cell),
            ])
        t = Table(data, colWidths=[20 * mm, 50 * mm, 55 * mm, 45 * mm], repeatRows=1)
        t.setStyle(TableStyle(styles))
        return t

    # ---- document assembly
    story = [PT(g["user"] + " - Prompting Guide", h1),
             RT(f"Updated {escape(g['updated'])} &nbsp;-&nbsp; source: {escape(g['source'])}", sub),
             Spacer(1, 6)]

    s = g["snapshot"]
    chips = Table([[f"EXCELLENT  {s['excellent']}", f"GOOD  {s['good']}", f"BAD  {s['bad']}",
                    f"REVIEWED  {s['reviewed']}"]], colWidths=[CONTENT_W / 4] * 4)
    chips.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (0, 0), HEX(BANDS["excellent"][1])),
        ("BACKGROUND", (1, 0), (1, 0), HEX(BANDS["good"][1])),
        ("BACKGROUND", (2, 0), (2, 0), HEX(BANDS["bad"][1])),
        ("BACKGROUND", (3, 0), (3, 0), HEX("#eef1f4")),
        ("TEXTCOLOR", (0, 0), (0, 0), HEX(BANDS["excellent"][0])),
        ("TEXTCOLOR", (1, 0), (1, 0), HEX(BANDS["good"][0])),
        ("TEXTCOLOR", (2, 0), (2, 0), HEX(BANDS["bad"][0])),
        ("FONTNAME", (0, 0), (-1, -1), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("TOPPADDING", (0, 0), (-1, -1), 8), ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
        ("GRID", (0, 0), (-1, -1), 3, colors.white),
    ]))
    story += [chips, Spacer(1, 6),
              boxed([RT(f"<b>Most common gap:</b> {escape(s['common_gap'])}", body), Spacer(1, 3),
                     RT(f"<b>Strongest habit:</b> {escape(s['strongest_habit'])}", body), Spacer(1, 3),
                     RT(f"<b>Trend:</b> {escape(s['trend'])}", body)], "#f7f8fa", "#dfe3e8"),
              Spacer(1, 5),
              boxed([RT("<b>How to read each prompt below.</b> The <b>scorecard</b> lists every rubric "
                       "dimension with a status - <font color='#1a7f37'><b>Met</b></font>, "
                       "<font color='#9a6700'><b>Partial</b></font>, "
                       "<font color='#b42318'><b>Missing</b></font>, or <b>n/a</b> (not needed by this task). "
                       "The score is the weighted roll-up of the applicable rows (Met = full, Partial = half, "
                       "Missing = 0). The <b>transformation table</b> turns each gap into a concrete rewrite "
                       "and the principle it teaches.", body)], "#eef3fb", "#b9cdf0"),
              Spacer(1, 10)]

    for section in g["sections"]:
        accent, light = BANDS[section["band"]]
        bar = Table([[RT(f"<b>{escape(section['title'])}</b>",
                         ParagraphStyle("sect", parent=body, fontSize=13,
                                        textColor=colors.white))]], colWidths=[CONTENT_W])
        bar.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, -1), HEX(accent)),
                                 ("LEFTPADDING", (0, 0), (-1, -1), 8),
                                 ("TOPPADDING", (0, 0), (-1, -1), 4),
                                 ("BOTTOMPADDING", (0, 0), (-1, -1), 4)]))
        story += [bar, Spacer(1, 6)]

        for ex in section["examples"]:
            head = Table([[RT(f"<font color='{accent}'><b>{escape(BAND_LABEL[section['band']])} "
                              f"&middot; score {ex['score']} &middot; {escape(ex.get('verdict', ''))}</b></font>"
                              f" &nbsp; <b>{escape(ex['label'])}</b>", body),
                           RT(f"<font color='#888888' size=8>{escape(ex['kind'])} &middot; "
                              f"{escape(ex['date'])}</font>",
                              ParagraphStyle("r", parent=body, alignment=2))]],
                         colWidths=[CONTENT_W * 0.72, CONTENT_W * 0.28])
            head.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), HEX(light)),
                ("BOX", (0, 0), (-1, -1), 0.75, HEX(accent)),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 8), ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 6), ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]))
            promptbox = boxed([Paragraph('"' + escape(ex["prompt"]) + '"', prompt)], "#ffffff", "#cfd4da")
            # keep the header glued to the prompt so a header never orphans at a page end
            story += [KeepTogether([head, Spacer(1, 3), promptbox]), Spacer(1, 5)]
            story += [RT(f"<b>Why it lands here.</b> {escape(ex['why'])}", body), Spacer(1, 2)]
            if ex.get("chain"):
                story += [RT(f"<b>Chain step.</b> {escape(ex['chain'])}", body), Spacer(1, 2)]
            story += [Spacer(1, 3), RT("<b>Rubric scorecard</b>", cellb), Spacer(1, 2),
                      rubric_table(ex), Spacer(1, 6)]
            tt = transform_table(ex)
            if tt is not None:
                story += [RT("<b>Transformation - turn each gap into a rewrite</b>", cellb),
                          Spacer(1, 2), tt, Spacer(1, 5)]
            else:
                story += [RT("<i>No changes needed - imitate this one.</i>",
                             ParagraphStyle("none", parent=body, textColor=HEX(accent))),
                          Spacer(1, 5)]
            if ex.get("fix"):
                story += [boxed([Paragraph("<b>Full rewritten prompt:</b> " + escape(ex["fix"]), fixstyle)],
                                "#eef3fb", "#b9cdf0"), Spacer(1, 4)]
            story += [Spacer(1, 8)]

    # habits
    hb = [RT("<b>Habits to build</b>", ParagraphStyle("hb", parent=body, fontSize=12, spaceAfter=4))]
    for it in g["habits_build"]:
        hb += [RT("&bull;&nbsp; " + escape(it), body), Spacer(1, 3)]
    story += [Spacer(1, 2), boxed(hb, "#f7f8fa", "#dfe3e8"), Spacer(1, 8)]

    hh = [RT("<b>Habits you already have</b>", ParagraphStyle("hh", parent=body, fontSize=12, spaceAfter=4))]
    for it in g["habits_have"]:
        hh += [RT("&bull;&nbsp; " + escape(it), body), Spacer(1, 3)]
    story += [boxed(hh, BANDS["excellent"][1], BANDS["excellent"][0])]

    SimpleDocTemplate(str(out), pagesize=A4, leftMargin=20 * mm, rightMargin=20 * mm,
                      topMargin=16 * mm, bottomMargin=16 * mm,
                      title=f"{g['user']} - Prompting Guide").build(story)


# -------------------------------------------------------------------------- DOCX
def render_docx(g: dict, out: Path) -> None:
    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Pt, RGBColor

    def hexrgb(h):
        return RGBColor.from_string(h.lstrip("#"))

    def shade_cell(c, fill):
        tcPr = c._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), fill.lstrip("#"))
        tcPr.append(shd)

    doc = docx.Document()

    def box(text, fill, italic=False, color=None, bold_lead=None):
        tbl = doc.add_table(rows=1, cols=1); tbl.style = "Table Grid"
        c = tbl.cell(0, 0); shade_cell(c, fill)
        p = c.paragraphs[0]
        if bold_lead:
            r = p.add_run(bold_lead + " "); r.bold = True
            if color: r.font.color.rgb = hexrgb(color)
        r = p.add_run(text); r.italic = italic
        if color: r.font.color.rgb = hexrgb(color)
        doc.add_paragraph()

    def small(c, text, bold=False, color=None, align=None):
        p = c.paragraphs[0]
        if align is not None: p.alignment = align
        r = p.add_run(text); r.font.size = Pt(8); r.bold = bold
        if color: r.font.color.rgb = hexrgb(color)

    title = doc.add_heading(f"{g['user']} - Prompting Guide", level=0)
    st = doc.add_paragraph(f"Updated {g['updated']}  -  source: {g['source']}")
    st.runs[0].font.size = Pt(9); st.runs[0].font.color.rgb = hexrgb("#666666")

    s = g["snapshot"]
    chip = doc.add_table(rows=1, cols=4); chip.style = "Table Grid"
    for i, (name, n, band) in enumerate([("EXCELLENT", s["excellent"], "excellent"),
                                         ("GOOD", s["good"], "good"), ("BAD", s["bad"], "bad"),
                                         ("REVIEWED", s["reviewed"], None)]):
        c = chip.cell(0, i); shade_cell(c, BANDS[band][1] if band else "#eef1f4")
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"{name}\n{n}"); r.bold = True; r.font.size = Pt(11)
        if band: r.font.color.rgb = hexrgb(BANDS[band][0])
    doc.add_paragraph()
    box(s["common_gap"], "#f7f8fa", bold_lead="Most common gap:")
    box(s["strongest_habit"], "#f7f8fa", bold_lead="Strongest habit:")
    box("Each prompt has a scorecard (Met / Partial / Missing / n-a per rubric dimension; the "
        "score is the weighted roll-up of applicable rows) and a transformation table (each gap "
        "-> a concrete rewrite + the principle it teaches).", "#eef3fb", color="#0b3d91",
        bold_lead="How to read this:")

    for section in g["sections"]:
        accent, light = BANDS[section["band"]]
        h = doc.add_heading(section["title"], level=1)
        for r in h.runs: r.font.color.rgb = hexrgb(accent)
        for ex in section["examples"]:
            hp = doc.add_paragraph()
            rb = hp.add_run(f"{BAND_LABEL[section['band']]} - score {ex['score']} - {ex.get('verdict','')}   ")
            rb.bold = True; rb.font.color.rgb = hexrgb(accent)
            rl = hp.add_run(ex["label"]); rl.bold = True
            rk = hp.add_run(f"   ({ex['kind']} - {ex['date']})")
            rk.font.size = Pt(8); rk.font.color.rgb = hexrgb("#888888")
            box('"' + ex["prompt"] + '"', "#ffffff", italic=True)
            box(ex["why"], light, bold_lead="Why it lands here.")
            if ex.get("chain"):
                box(ex["chain"], light, bold_lead="Chain step.")

            # rubric scorecard table
            doc.add_paragraph().add_run("Rubric scorecard").bold = True
            by = dim_row(ex["rubric"])
            rt = doc.add_table(rows=1, cols=3); rt.style = "Table Grid"
            hdr = rt.rows[0].cells
            for c, txt in zip(hdr, ("Rubric dimension", "Status", "Why (evidence)")):
                shade_cell(c, "#334155"); small(c, txt, bold=True, color="#ffffff")
            for did, name, _dl, _w in DIMS:
                row = by.get(did, {"verdict": "na", "evidence": ""})
                vlabel, acc, lt = VERDICT[row["verdict"]]
                cells = rt.add_row().cells
                small(cells[0], f"{did}  {name}")
                shade_cell(cells[1], acc); small(cells[1], vlabel, bold=True, color="#ffffff",
                                                 align=WD_ALIGN_PARAGRAPH.CENTER)
                small(cells[2], row.get("evidence", ""))
            doc.add_paragraph()

            # transformation table
            rows = ex.get("transformation") or []
            if rows:
                doc.add_paragraph().add_run("Transformation - turn each gap into a rewrite").bold = True
                tt = doc.add_table(rows=1, cols=4); tt.style = "Table Grid"
                for c, txt in zip(tt.rows[0].cells,
                                  ("Rubric", "You wrote", "Best-practice rewrite", "Principle")):
                    shade_cell(c, "#0b3d91"); small(c, txt, bold=True, color="#ffffff")
                for row in rows:
                    cells = tt.add_row().cells
                    small(cells[0], row.get("id", ""), bold=True)
                    small(cells[1], row.get("before", ""))
                    small(cells[2], row.get("after", ""), color="#0b3d91")
                    small(cells[3], row.get("principle", ""))
                doc.add_paragraph()
            else:
                p = doc.add_paragraph(); r = p.add_run("No changes needed - imitate this one.")
                r.italic = True; r.font.color.rgb = hexrgb(accent)

            if ex.get("fix"):
                box(ex["fix"], "#eef3fb", color="#0b3d91", bold_lead="Full rewritten prompt:")

    for title_txt, items, accent in [("Habits to build", g["habits_build"], "#111111"),
                                      ("Habits you already have", g["habits_have"], BANDS["excellent"][0])]:
        h = doc.add_heading(title_txt, level=1)
        for r in h.runs: r.font.color.rgb = hexrgb(accent)
        for it in items:
            doc.add_paragraph(it, style="List Bullet")

    doc.save(str(out))


# ---------------------------------------------------------------------------- MD
def _mdcell(text: str) -> str:
    return str(text).replace("|", "\\|").replace("\n", " ").strip()


def render_md(g: dict, out: Path) -> None:
    L = []
    L.append(f"# Prompting Guide - {g['user']}")
    L.append("")
    L.append(f"_Last updated {g['updated']}. Built from prompt-critic reviews of the journal logs._")
    L.append("")
    s = g["snapshot"]
    L.append("## Snapshot")
    L.append(f"- Prompts reviewed: {s['reviewed']} (excellent {s['excellent']} / good "
             f"{s['good']} / bad {s['bad']}) - source: `{g['source']}`")
    L.append(f"- Trend vs. last update: {s['trend']}")
    L.append(f"- Most common gap: **{_mdcell(s['common_gap'])}**")
    L.append(f"- Strongest habit: **{_mdcell(s['strongest_habit'])}**")
    cov = s.get("coverage")
    if cov:
        projects = ", ".join(cov.get("projects", [])) or "-"
        span = f"{cov.get('from', '?')} → {cov.get('to', '?')}"
        L.append(f"- Coverage: {cov.get('files', '?')} files across projects [{_mdcell(projects)}] ({span})")
    L.append("")
    L.append("_Each prompt below has a **rubric scorecard** (Met / Partial / Missing / n-a per "
             "dimension; the score is the weighted roll-up of the applicable rows) and a "
             "**transformation table** (each gap -> a concrete rewrite + the principle it teaches)._")
    L.append("")

    for section in g["sections"]:
        L.append(f"## {section['title']}")
        L.append("")
        for ex in section["examples"]:
            L.append(f"### {ex['label']} - band: {section['band']}, score {ex['score']} "
                     f"({ex.get('verdict','')})")
            L.append(f"> {ex['prompt']}")
            L.append("")
            L.append(f"_Source: `{g['source']}` &middot; {ex['date']} &middot; {ex['kind']}_")
            L.append("")
            L.append(f"- **Why it lands here:** {_mdcell(ex['why'])}")
            if ex.get("chain"):
                L.append(f"- **Chain step:** {_mdcell(ex['chain'])}")
            L.append("")
            # scorecard
            L.append("**Rubric scorecard**")
            L.append("")
            L.append("| Rubric dimension | Status | Why (evidence) |")
            L.append("|---|---|---|")
            by = dim_row(ex["rubric"])
            for did, name, _dl, _w in DIMS:
                row = by.get(did, {"verdict": "na", "evidence": ""})
                L.append(f"| {did} {name} | {MD_MARK[row['verdict']]} | "
                         f"{_mdcell(row.get('evidence',''))} |")
            L.append("")
            # transformation
            rows = ex.get("transformation") or []
            if rows:
                L.append("**Transformation - turn each gap into a rewrite**")
                L.append("")
                L.append("| Rubric | You wrote | Best-practice rewrite | Principle |")
                L.append("|---|---|---|---|")
                for row in rows:
                    L.append(f"| {row.get('id','')} | {_mdcell(row.get('before',''))} | "
                             f"{_mdcell(row.get('after',''))} | {_mdcell(row.get('principle',''))} |")
                L.append("")
            else:
                L.append("_No changes needed - imitate this one._")
                L.append("")
            if ex.get("fix"):
                L.append(f"- **Full rewritten prompt:** _{_mdcell(ex['fix'])}_")
                L.append("")

    L.append("## Habits to build")
    for it in g["habits_build"]:
        L.append(f"- {it}")
    L.append("")
    L.append("## Habits you already have")
    for it in g["habits_have"]:
        L.append(f"- {it}")
    L.append("")
    out.write_text("\n".join(L), encoding="utf-8")


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("json", help="path to guides/<user>.json")
    ap.add_argument("--pdf", action="store_true")
    ap.add_argument("--docx", action="store_true")
    ap.add_argument("--md", action="store_true")
    args = ap.parse_args()

    src = Path(args.json)
    g = load(src)

    for w in verify(g):
        print("WARNING: rubric roll-up disagrees with stored score:")
        print(w)

    both = not (args.pdf or args.docx or args.md)
    done = []
    if args.pdf or both:
        out = src.with_suffix(".pdf"); render_pdf(g, out); done.append(str(out))
    if args.docx or both:
        out = src.with_suffix(".docx"); render_docx(g, out); done.append(str(out))
    if args.md or both:
        out = src.with_suffix(".md"); render_md(g, out); done.append(str(out))
    print("Wrote:\n  " + "\n  ".join(done))
    return 0


if __name__ == "__main__":
    sys.exit(main())
